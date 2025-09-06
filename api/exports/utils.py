from __future__ import annotations

from io import BytesIO
import re
import zipfile

from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import html


def _escape_text(s: str) -> str:
    # Basic HTML escape to avoid accidental HTML rendering in downstream tools
    return html.escape(s, quote=False)


def proposal_json_to_markdown(proposal: dict) -> str:
    meta = proposal.get('meta', {})
    sections = proposal.get('sections', {}) or {}
    lines = []
    title = _escape_text(meta.get('title') or 'Proposal')
    lines.append(f"# {title}")
    for key, section in sections.items():
        title = _escape_text(section.get('title') or key)
        content = _escape_text(str(section.get('content') or ''))
        lines.append("")
        lines.append(f"## {title}")
        lines.append(content)
    return "\n".join(lines)


def _normalize_pdf_for_checksum(data: bytes) -> bytes:
    # Remove/normalize non-deterministic parts: document ID and xref offset
    try:
        data = re.sub(rb"/ID\s*\[\s*<[^>]*>\s*<[^>]*>\s*\]", b"/ID [<000000><000000>]", data)
        data = re.sub(rb"startxref\s*\d+", b"startxref 0", data)
        data = re.sub(rb"/CreationDate\s*\(D:[^\)]+\)", b"/CreationDate (D:19700101000000Z)", data)
        data = re.sub(rb"/ModDate\s*\(D:[^\)]+\)", b"/ModDate (D:19700101000000Z)", data)
    except Exception:
        pass
    return data


def render_pdf_from_text(text: str) -> tuple[bytes, str]:
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=LETTER)
    # Deterministic metadata (guarded: ReportLab internals may differ across versions)
    try:  # Accessing protected internals for deterministic metadata; ignore type
        info = c._doc.info  # type: ignore[attr-defined]
        info.title = "Granterstellar Export"
        info.author = "Granterstellar"
        info.creator = "Granterstellar"
        info.producer = "ReportLab"
        info.creationDate = 'D:19700101000000Z'
        info.modDate = 'D:19700101000000Z'
    except Exception:  # pragma: no cover - defensive
        pass
    width, height = LETTER
    y = height - 72
    for line in text.splitlines():
        if y < 72:
            c.showPage()
            y = height - 72
        c.drawString(72, y, line[:1000])
        y -= 14
    c.showPage()
    c.save()
    data = buffer.getvalue()
    # Post-process raw PDF to enforce epoch Creation/Mod dates so raw bytes deterministic
    try:
        data = re.sub(rb"/CreationDate\s*\(D:[^\)]+\)", b"/CreationDate (D:19700101000000Z)", data)
        data = re.sub(rb"/ModDate\s*\(D:[^\)]+\)", b"/ModDate (D:19700101000000Z)", data)
    except Exception:  # pragma: no cover - defensive
        pass
    normalized = _normalize_pdf_for_checksum(data)
    # Defer to common checksum utility for consistency
    try:
        from app.common.files import compute_checksum  # local import to avoid early app loading side-effects
        checksum = compute_checksum(normalized).hex
    except Exception:
        # Fallback to direct hashlib if utility unavailable (defensive during migrations)
        import hashlib as _hl
        checksum = _hl.sha256(normalized).hexdigest()
    return data, checksum


def render_docx_from_markdown(md: str) -> tuple[bytes, str]:
    """
    Minimal Markdown -> DOCX renderer.
    Supports:
      - # H1 and ## H2 headings
      - Paragraphs
      - Blank lines for spacing
    This keeps the canonical markdown as the source of truth and does a light render.
    """
    doc = Document()
    # Set a readable base style size
    try:  # Styles may not always be mutable
        style = doc.styles['Normal']  # type: ignore[index]
        style.font.size = Pt(11)  # type: ignore[attr-defined]
    except Exception:  # pragma: no cover - defensive
        pass

    for raw in md.splitlines():
        line = raw.rstrip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith('# '):
            p = doc.add_heading(line[2:].strip(), level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue
        if line.startswith('## '):
            p = doc.add_heading(line[3:].strip(), level=2)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue
        doc.add_paragraph(line)

    # Deterministic core properties
    try:
        core = doc.core_properties
        core.title = "Granterstellar Export"
        core.author = "Granterstellar"
        from datetime import datetime, timezone
        epoch = datetime(1970, 1, 1, tzinfo=timezone.utc)
        core.created = epoch
        core.modified = epoch
        core.last_printed = epoch
    except Exception:
        pass

    bio = BytesIO()
    doc.save(bio)
    raw = bio.getvalue()

    # Normalize DOCX zip for determinism: fixed timestamps and sorted entries
    def _normalize_docx_zip(data: bytes) -> bytes:
        src = BytesIO(data)
        out_bio = BytesIO()
        with zipfile.ZipFile(src, 'r') as zin, zipfile.ZipFile(out_bio, 'w', compression=zipfile.ZIP_DEFLATED) as zout:
            for name in sorted(zin.namelist()):
                content = zin.read(name)
                zi = zipfile.ZipInfo(filename=name, date_time=(1980, 1, 1, 0, 0, 0))
                zi.compress_type = zipfile.ZIP_DEFLATED
                # Ensure consistent permissions/external attrs
                zi.external_attr = 0o600 << 16
                zout.writestr(zi, content)
        return out_bio.getvalue()

    data = _normalize_docx_zip(raw)
    try:
        from app.common.files import compute_checksum
        checksum = compute_checksum(data).hex
    except Exception:
        import hashlib as _hl
        checksum = _hl.sha256(data).hexdigest()
    return data, checksum
